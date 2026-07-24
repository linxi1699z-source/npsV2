from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "NPS管理后台需求文档.docx"
OUTPUT = ROOT / "docs" / "NPS管理后台需求文档-最新原型.docx"
ASSET_DIR = ROOT / "docs" / "assets" / "latest-prd"
PROTO_URL = "http://localhost:4174/prototypes/nps-task/index.html"

BLUE = "1677FF"
LIGHT_BLUE = "E6F4FF"
HEADER_FILL = "F0F5FF"
LIGHT_GRAY = "FAFAFA"
BORDER = "D9D9D9"
TEXT = "262626"
MUTED = "595959"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_fixed(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    requested = [round(width * 1440) for width in widths]
    target_total = 10268
    scale = target_total / sum(requested)
    width_dxa = [round(value * scale) for value in requested]
    width_dxa[-1] += target_total - sum(width_dxa)
    total_dxa = sum(width_dxa)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "100")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for value in width_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(value))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, value in enumerate(width_dxa):
            row.cells[idx].width = Inches(value / 1440)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=9, bold=False, color=TEXT):
    run.font.name = "Arial Unicode MS"
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), "Arial Unicode MS")
    r_fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(str(header)), size=font_size, bold=True)
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if idx in (2, 3, 4) and len(headers) == 6:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(str(value)), size=font_size)
    set_table_fixed(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), size=10.5, bold=True)
        style_run(p.add_run(text[len(bold_prefix):]), size=10.5)
    else:
        style_run(p.add_run(text), size=10.5)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        style_run(p.add_run(item), size=10.5)


def clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    heading_settings = {
        "Heading 1": (16, BLUE, 14, 7),
        "Heading 2": (13, BLUE, 11, 5),
        "Heading 3": (11.5, "1F1F1F", 9, 4),
        "Heading 4": (10.5, "434343", 7, 3),
    }
    for name, (size, color, before, after) in heading_settings.items():
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def font_path():
    for candidate in (
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ):
        if Path(candidate).exists():
            return candidate
    return None


def draw_flow(path, title, columns):
    width, height = 1500, 500
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 34) if fp else ImageFont.load_default()
    box_font = ImageFont.truetype(fp, 22) if fp else ImageFont.load_default()
    note_font = ImageFont.truetype(fp, 18) if fp else ImageFont.load_default()
    draw.text((48, 32), title, font=title_font, fill="#1F1F1F")
    box_w = 220
    box_h = 86
    gap = 56
    start_x = 42
    y = 165
    for index, (label, note) in enumerate(columns):
        x = start_x + index * (box_w + gap)
        fill = "#E6F4FF" if index not in (0, len(columns) - 1) else "#F6FFED"
        outline = "#1677FF" if index not in (0, len(columns) - 1) else "#52C41A"
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=12, fill=fill, outline=outline, width=3)
        bbox = draw.textbbox((0, 0), label, font=box_font)
        tx = x + (box_w - (bbox[2] - bbox[0])) / 2
        draw.text((tx, y + 17), label, font=box_font, fill="#262626")
        if note:
            bbox2 = draw.textbbox((0, 0), note, font=note_font)
            tx2 = x + (box_w - (bbox2[2] - bbox2[0])) / 2
            draw.text((tx2, y + 52), note, font=note_font, fill="#595959")
        if index < len(columns) - 1:
            ax1 = x + box_w + 10
            ax2 = x + box_w + gap - 10
            ay = y + box_h / 2
            draw.line((ax1, ay, ax2, ay), fill="#8C8C8C", width=3)
            draw.polygon([(ax2, ay), (ax2 - 13, ay - 8), (ax2 - 13, ay + 8)], fill="#8C8C8C")
    draw.rounded_rectangle((42, 330, 1458, 430), radius=12, fill="#FAFAFA", outline="#D9D9D9", width=2)
    note = "异常与限制：表单校验失败时停留当前页面；状态或关联关系不满足操作条件时不修改数据，并展示明确提示。"
    draw.text((72, 365), note, font=note_font, fill="#595959")
    image.save(path)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(7.0))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    style_run(c.add_run(caption), size=9, color=MUTED)


def h(doc, level, text):
    return doc.add_heading(text, level=level)


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    q_flow = ASSET_DIR / "questionnaire-flow.png"
    t_flow = ASSET_DIR / "delivery-task-flow.png"
    draw_flow(q_flow, "问卷配置主流程", [
        ("进入问卷列表", "查询 / 新增"),
        ("设置基础信息", "按类型联动"),
        ("设置问卷问题", "问题 / 分页"),
        ("提交问卷", "状态生效中"),
        ("后续维护", "多语言 / 状态"),
    ])
    draw_flow(t_flow, "投放任务主流程", [
        ("进入投放任务", "查询 / 新增"),
        ("选择功能模块", "全局需选类型"),
        ("关联问卷", "仅生效中"),
        ("配置投放", "时间 / 版本 / 人群"),
        ("任务执行", "待投放至完成"),
    ])

    doc = Document(SOURCE)
    clear_document_body(doc)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    style_run(title.add_run("NPS管理后台需求文档"), size=22, bold=True, color="1F1F1F")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    style_run(subtitle.add_run("最新原型对应版本"), size=11, color=MUTED)

    h(doc, 1, "版本信息")
    add_body(doc, "版本号：V2.0")
    add_body(doc, "创建日期：2026-07-22")
    add_body(doc, "作者：谢敏")
    h(doc, 2, "变更日志")
    add_table(doc, ["时间", "版本号", "变更人", "主要变更内容"], [[
        "2026-07-22", "V2.0", "谢敏",
        "按最新可见原型重构需求范围，覆盖问卷列表、问卷新增/编辑、设置问卷问题、多语言配置、投放任务和投放人群。"
    ]], [1.05, 0.75, 0.85, 4.55])

    h(doc, 1, "文档说明")
    h(doc, 2, "名词解释")
    terms = [
        ("NPS", "Net Promoter Score，用于衡量用户推荐意愿和满意度的指标。"),
        ("问卷", "由基础信息、问题、分页和多语言文案组成，可被投放任务关联。"),
        ("弹窗问卷(全局)", "全局功能使用的弹窗问卷，基础信息中包含两个固定类型的全局问题。"),
        ("弹窗问卷(用研/设计)", "用于用户研究或设计反馈的弹窗问卷，基础信息中包含富文本弹窗文案。"),
        ("功能问卷", "用于非全局功能模块投放的普通问卷。"),
        ("生效中 / 已禁用", "问卷列表展示状态。生效中问卷可被新任务关联；已禁用问卷不可被新任务关联。"),
        ("分页", "问卷问题的页面分组。第一页默认存在，新增分页后问题在下一页展示。"),
        ("多语言文件", "按当前问卷文案生成并回传的 XLSX 文件，用于维护多语言内容。"),
        ("投放任务", "选择问卷并配置功能模块、投放时间、投放版本及投放人群的执行单元。"),
        ("提交用户数", "提交部分问卷或提交完整问卷的去重用户数。"),
        ("曝光状态", "用户维度状态，包括未曝光、已曝光。"),
        ("投放状态", "用户维度状态，包括未开始、已关闭、已提交部分问卷、已提交全部问卷。"),
    ]
    add_table(doc, ["术语 / 缩略词", "说明"], terms, [1.75, 5.45], 9)

    h(doc, 1, "需求背景")
    h(doc, 2, "产品 / 数据现状调研")
    add_body(doc, "NPS 管理后台需要提供统一的问卷配置和投放任务管理能力。产品或运营人员可维护问卷结构、问题、分页及多语言文案，再创建投放任务，将生效中的问卷投放给指定用户。当前原型的可见导航包括问卷列表和投放任务，并提供投放人群明细查询。")
    h(doc, 2, "用户调研")
    add_body(doc, "待补充。目标使用者包括产品、运营、用户研究及相关后台管理员。")
    h(doc, 2, "竞品分析")
    add_body(doc, "待补充。")

    h(doc, 1, "目标与关键指标（KPI）")
    add_body(doc, "目标：", bold_prefix="目标：")
    add_bullets(doc, [
        "支持问卷从创建、编辑、问题配置、多语言维护到启用、禁用和删除的完整后台管理流程。",
        "支持投放任务按功能模块和问卷类型准确关联问卷，并通过默认投放版本降低配置成本。",
        "支持按任务查看投放人群的曝光和提交状态，便于运营核对执行结果。",
        "通过明确的字段校验、状态限制和确认弹窗降低误操作风险。",
    ], numbered=True)
    add_body(doc, "KPI：", bold_prefix="KPI：")
    add_bullets(doc, [
        "问卷及投放任务表单配置成功率达到 99% 以上。",
        "因字段遗漏、问卷类型联动错误导致的提交失败率低于 1%。",
        "列表查询、状态更新和页面跳转在正常网络条件下 2 秒内完成。",
        "重复提交不产生重复问卷或重复任务记录。",
    ], numbered=True)

    h(doc, 1, "功能详细说明")
    h(doc, 2, "5.1 主要功能范围")
    scope = [
        ("问卷列表", "查询、分页、查看状态、新增、编辑、启用、禁用、删除、多语言入口", "默认按问卷 ID 倒序展示。"),
        ("新增 / 编辑问卷", "设置问卷名称、网页显示名、问卷类型及类型对应基础信息", "编辑时问卷类型不可修改。"),
        ("设置问卷问题", "支持评分、单选题、多选题、开放题、纯文本；支持新增问题和新增分页", "第一页无分页标题；新增分页可折叠、展开和删除。"),
        ("多语言配置", "下载当前问卷文案模板、上传 XLSX、语言切换预览、提交", "未上传时为新增模式；已上传时为编辑模式。"),
        ("投放任务", "查询、列表、新增、编辑、导出、删除", "仅待投放任务可编辑和删除；仅投放完成可导出。"),
        ("投放人群", "按 RingConn ID、投放状态、曝光状态查询用户明细", "从任务列表的投放人群入口进入。"),
    ]
    add_table(doc, ["模块", "功能点", "备注"], scope, [1.35, 3.6, 2.25], 8.7)

    h(doc, 2, "5.2 产品流程图")
    add_figure(doc, q_flow, "图 1  问卷配置主流程")
    add_figure(doc, t_flow, "图 2  投放任务主流程")

    h(doc, 2, "5.3 交互原型")
    p = doc.add_paragraph()
    style_run(p.add_run("原型总入口："), size=10.5, bold=True)
    add_hyperlink(p, PROTO_URL, PROTO_URL)
    p2 = doc.add_paragraph()
    style_run(p2.add_run("问卷列表："), size=10.5, bold=True)
    add_hyperlink(p2, f"{PROTO_URL}#survey-list", f"{PROTO_URL}#survey-list")
    p3 = doc.add_paragraph()
    style_run(p3.add_run("投放任务："), size=10.5, bold=True)
    add_hyperlink(p3, f"{PROTO_URL}#nps-task", f"{PROTO_URL}#nps-task")

    h(doc, 2, "5.4 功能说明")
    h(doc, 3, "5.4.1 导航与页面范围")
    add_body(doc, "一级导航为 NPS管理；可见二级导航为问卷列表、投放任务。点击二级导航后更新选中状态、面包屑和主内容区域。各新增、编辑、配置及明细页面通过面包屑体现当前路径。")
    add_table(doc, ["入口", "目标页面", "面包屑", "说明"], [
        ("NPS管理 / 问卷列表", "问卷列表", "NPS管理 / 问卷列表", "系统默认入口。"),
        ("+新增问卷", "新增问卷", "NPS管理 / 问卷列表 / 新增问卷", "进入两步式表单。"),
        ("编辑问卷", "编辑问卷", "NPS管理 / 问卷列表 / 编辑问卷", "回显数据；问卷类型锁定。"),
        ("新增/编辑多语言", "多语言配置", "NPS管理 / 问卷列表 / 新增或编辑多语言", "按文件状态进入对应模式。"),
        ("NPS管理 / 投放任务", "投放任务", "NPS管理 / 投放任务", "任务查询与维护。"),
        ("投放人群", "投放人群", "NPS管理 / 投放任务 / 投放人群", "展示指定任务用户明细。"),
    ], [1.55, 1.4, 2.45, 1.8], 8.5)

    h(doc, 3, "5.4.2 问卷列表")
    add_body(doc, "页面目标：查询并维护问卷；列表默认按问卷 ID 倒序展示，后创建问卷显示在前。")
    h(doc, 4, "搜索条件")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("questionnaire_name", "问卷名称", "否", "string", "空", "输入框；按问卷名称模糊匹配。"),
        ("questionnaire_status", "状态", "否", "enum", "所有状态", "单选下拉：所有状态、生效中、已禁用。"),
        ("questionnaire_type", "问卷类型", "否", "enum", "所有问卷类型", "单选下拉：所有问卷类型、弹窗问卷(全局)、弹窗问卷(用研/设计)、功能问卷。"),
    ], [1.25, 1.05, 0.65, 0.75, 0.9, 2.6], 8)
    h(doc, 4, "列表字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("questionnaire_id", "问卷 ID", "是", "int", "系统生成", "唯一标识；列表倒序。"),
        ("questionnaire_name", "问卷名称", "是", "string", "无", "展示后台管理名称。"),
        ("status", "状态", "是", "enum", "生效中", "状态标签：生效中、已禁用。"),
        ("questionnaire_type", "问卷类型", "是", "enum", "无", "弹窗问卷(全局)、弹窗问卷(用研/设计)、功能问卷。"),
        ("i18n_file_status", "多语言文件", "是", "enum", "未上传", "非全局问卷展示未上传/已上传；弹窗问卷(全局)展示 --。"),
        ("creator", "创建者", "是", "string", "当前用户", "展示创建人。"),
        ("updated_at", "最近一次更新时间", "是", "datetime", "系统时间", "格式 YYYY-MM-DD hh:mm；提交、状态变更、删除前更新。"),
        ("operation", "操作", "是", "button[]", "无", "配置多语言、新增/编辑多语言、编辑问卷、启用/禁用、删除；按钮按 Ant Design 语义色展示。"),
    ], [1.25, 1.05, 0.65, 0.75, 0.9, 2.6], 8)
    h(doc, 4, "页面交互")
    add_table(doc, ["触发位置", "触发动作", "前置条件", "系统处理", "结果 / 提示"], [
        ("查询", "点击", "条件可为空", "按名称、状态、类型组合筛选", "刷新列表和总数。"),
        ("+新增问卷", "点击", "有新增权限", "打开新增问卷页面", "进入设置基础信息。"),
        ("编辑问卷", "点击", "问卷存在", "回显问卷及问题数据；锁定问卷类型", "提交后更新原记录。"),
        ("禁用", "点击", "状态=生效中", "打开禁用确认弹窗", "确认后状态=已禁用。"),
        ("启用", "点击", "状态=已禁用", "打开启用确认弹窗", "确认后状态=生效中。"),
        ("删除", "点击", "问卷未关联投放任务", "打开删除确认弹窗", "确认后从列表及可选问卷池移除。"),
    ], [1.15, 0.75, 1.55, 2.15, 1.6], 8.2)

    h(doc, 3, "5.4.3 新增 / 编辑问卷 - 设置基础信息")
    add_body(doc, "页面采用步骤式表单。步骤页签仅展示进度，不允许点击切换；用户只能通过下一步、返回上一步、提交和取消完成导航。")
    h(doc, 4, "基础字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("questionnaire_name", "问卷名称", "是", "string", "空", "后台管理名称；为空提示：模板名称不能为空。"),
        ("web_display_name", "问卷名称(网页显示名)", "否", "string", "空", "问卷网页展示名称；编辑旧数据时默认回显问卷名称。"),
        ("questionnaire_type", "问卷类型", "是", "enum", "请选择", "单选下拉：弹窗问卷(全局)、弹窗问卷(用研/设计)、功能问卷；新增允许选择，编辑禁用。"),
        ("banner_title", "Banner标题", "是", "string", "空", "备注：APP通知中展示的标题。"),
        ("banner_subtitle", "Banner副标题", "是", "string", "空", "APP 通知副标题。"),
        ("banner_button", "Banner按钮文案", "是", "string", "空", "备注：APP通知中对应按钮文字。"),
        ("popup_copy", "弹窗文案", "条件必填", "richtext", "空", "仅弹窗问卷(用研/设计)展示；支持加粗、下划线、红/蓝/黑字体，最多 1000 字。"),
    ], [1.25, 1.1, 0.75, 0.8, 0.75, 2.55], 8)
    h(doc, 4, "问卷类型联动")
    add_table(doc, ["问卷类型", "基础信息附加内容", "下一步", "限制"], [
        ("弹窗问卷(全局)", "Banner 三字段 + 全局问题", "进入设置问卷问题", "全局问题固定为评分(全局)+多选题；只能调整文案和选项。"),
        ("弹窗问卷(用研/设计)", "Banner 三字段 + 富文本弹窗文案", "进入设置问卷问题", "弹窗文案最多 1000 字。"),
        ("功能问卷", "Banner 三字段", "进入设置问卷问题", "不展示全局问题和弹窗文案。"),
    ], [1.6, 2.35, 1.35, 1.9], 8.5)
    h(doc, 4, "弹窗问卷(全局)固定问题")
    add_table(doc, ["问题", "字段", "是否必填", "默认值", "规则 / 交互说明"], [
        ("评分(全局)", "问题类型", "是", "评分(全局)", "固定，不允许修改。"),
        ("评分(全局)", "问题标题", "是", "您有多大可能会向朋友或者同时推荐RingConn?", "允许修改文案。"),
        ("评分(全局)", "问题标题备注", "否", "空", "允许修改。"),
        ("评分(全局)", "选项", "是", "10", "可选择 5 或 10。"),
        ("评分(全局)", "最小值描述 / 最大值描述", "否", "不推荐 / 非常推荐", "允许修改。"),
        ("评分(全局)", "低 / 高 / 其他评分引导文案", "否", "预置文案", "允许修改。"),
        ("多选题", "问题类型", "是", "多选题", "固定，不允许修改。"),
        ("多选题", "选项", "是", "10 个预置选项", "允许修改、添加和删除选项；不展示问题标题和标题备注。"),
    ], [1.2, 1.55, 0.8, 2.0, 1.65], 8.2)
    add_body(doc, "多选题默认选项：戒指外观设计与质感、佩戴舒适度与尺寸、续航与充电体验、戒指质量与耐用度、App使用流畅度、App界面易读性、核心功能价值、数据监测准确性、数据解读帮助性、售后体验。")

    h(doc, 3, "5.4.4 设置问卷问题")
    h(doc, 4, "分页与问题容器")
    add_table(doc, ["对象", "默认状态", "操作", "规则 / 交互说明"], [
        ("问卷第一页", "默认展开", "新增问题", "不展示分页标题、折叠/展开、删除分页及分页外框。"),
        ("新增分页", "默认展开", "折叠 / 展开 / 删除分页", "按钮下方新增“问卷第 N 页”；标题栏使用浅蓝色，分页边框为虚线。删除分页时同步删除页内全部问题。"),
        ("问题", "每页默认 1 个评分问题", "新增 / 删除", "问题使用实线卡片；问题之间保持间距；提交时每页至少保留 1 个问题。"),
        ("新增问题 / 新增分页", "可见", "点击", "两个按钮之间使用虚线分隔，并与分隔线保持不少于 20px 间距。"),
    ], [1.3, 1.3, 1.55, 3.05], 8.5)
    h(doc, 4, "通用问题字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("question_type", "问题类型", "是", "enum", "评分", "选项：评分、单选题、多选题、开放题、纯文本；变更后仅刷新当前问题字段。"),
        ("question_title", "问题标题", "是", "string", "空", "每种问题类型均需填写；纯文本类型用于展示正文标题。"),
        ("question_title_remark", "问题标题备注", "否", "string", "空", "展示在问题标题下方。"),
    ], [1.25, 1.1, 0.7, 0.75, 0.8, 2.6], 8)
    h(doc, 4, "各问题类型字段")
    add_table(doc, ["问题类型", "附加字段", "是否必填", "默认值", "规则 / 交互说明"], [
        ("评分", "选项", "是", "10", "单选下拉，选项 5、10；同步展示 1 至 N 的评分预览。"),
        ("评分", "最小值描述 / 最大值描述", "否", "空", "分别展示在评分两端。"),
        ("单选题", "选项", "是", "选项", "至少 1 项；使用单选圆形图标；支持添加、删除和其他选项。"),
        ("多选题", "选项", "是", "选项", "至少 1 项；使用复选方形图标；支持添加、删除和其他选项。"),
        ("单选题 / 多选题", "其他选项描述", "否", "空", "添加其他后展示；字段与问题字段对齐。"),
        ("单选题 / 多选题", "其他字符限制", "是", "500", "添加其他后展示；选项 500、1000、2000。"),
        ("开放题", "必填 / 选填", "是", "必填", "单选；选填时 H5 问题标题后自动追加“选填”。"),
        ("开放题", "字符限制", "是", "500", "单选下拉：500、1000、20000。"),
        ("纯文本", "无", "-", "-", "仅展示问题标题和问题标题备注，不展示答案控件。"),
    ], [1.45, 1.65, 0.75, 1.1, 2.25], 8.2)
    h(doc, 4, "选项交互")
    add_table(doc, ["触发位置", "触发动作", "前置条件", "系统处理", "结果 / 提示"], [
        ("添加选项", "点击", "单选题或多选题", "新增普通选项并放在其他选项之前", "默认文案为“选项”。"),
        ("添加其他", "点击", "当前问题无其他选项", "新增固定在末尾的其他选项", "文本框禁用，默认文案“其他”；按钮隐藏。"),
        ("删除其他", "点击", "当前问题存在其他选项", "删除该选项及其描述/字符限制", "重新展示添加其他按钮。"),
        ("删除问题", "点击", "问题存在", "删除当前问题", "若某页无问题，提交时提示完善问题。"),
    ], [1.15, 0.75, 1.55, 2.1, 1.65], 8.2)

    h(doc, 3, "5.4.5 多语言配置")
    add_body(doc, "入口位于问卷列表操作列。弹窗问卷(全局)的多语言文件显示 --，配置多语言按钮禁用；其他问卷按文件状态进入新增或编辑多语言页面。")
    h(doc, 4, "状态与页面")
    add_table(doc, ["多语言文件", "操作按钮", "页面内容", "提交结果"], [
        ("未上传", "新增多语言", "仅展示下载和上传步骤，不展示右侧预览", "上传后点击提交，状态更新为已上传并返回列表。"),
        ("已上传", "编辑多语言", "展示下载、上传及右侧语言切换+iPhone H5 预览", "重新上传并提交后更新文案并返回列表。"),
        ("--", "配置多语言（禁用）", "不可进入", "适用于弹窗问卷(全局)。"),
    ], [1.25, 1.45, 2.8, 1.7], 8.5)
    h(doc, 4, "下载与上传规则")
    add_table(doc, ["功能", "字段 / 文件", "规则 / 交互说明"], [
        ("下载文案", "{问卷名称}多语言模板.xlsx", "导出当前问卷全部可翻译文案；点击后加入浏览器下载队列。"),
        ("模板列", "A列 key；B列 中文名；后续每种语言一列", "语言：中文、繁体中文、英文、日本语、德语、西班牙语、法语、阿拉伯语、捷克语、土耳其语、波兰语、芬兰语、意大利语、匈牙利语、罗马尼亚语、韩语、立陶宛语、马来语。"),
        ("上传文件", ".xlsx", "只允许单个 XLSX；文件不超过 5M；选择文件后自动上传并展示 loading。"),
        ("上传成功", "文件解析完成", "新增模式提示“文件已上传，请提交更新。”；编辑模式提示“文件已上传，切换语言查看效果。”"),
        ("提交", "已上传文件", "未选择文件时提示“请先上传多语言文件。”；成功提示“多语言更新成功。”并返回列表。"),
        ("取消", "无", "不保存本次修改，返回问卷列表。"),
    ], [1.15, 2.1, 3.95], 8.5)

    h(doc, 3, "5.4.6 问卷状态与删除")
    h(doc, 4, "确认弹窗")
    add_table(doc, ["弹窗", "触发条件", "文案", "按钮", "确定逻辑", "取消 / 关闭"], [
        ("禁用确认", "状态=生效中，点击禁用", "确定需要将模板禁用吗？", "确定 / 取消", "状态更新为已禁用并刷新列表。", "关闭弹窗，数据不变。"),
        ("启用确认", "状态=已禁用，点击启用", "确定需要将模板启动吗？", "确定 / 取消", "状态更新为生效中并刷新列表。", "关闭弹窗，数据不变。"),
        ("删除确认", "未关联投放任务，点击删除", "删除后将无法继续使用该问卷，确定要继续吗？", "确定 / 取消", "逻辑删除问卷，刷新列表和关联问卷选项。", "关闭弹窗，数据不变。"),
    ], [0.9, 1.2, 1.65, 0.8, 1.55, 1.1], 7.7)
    h(doc, 4, "状态流转")
    add_table(doc, ["当前状态", "触发动作", "前置条件", "目标状态", "说明"], [
        ("生效中", "禁用", "当前状态仍为生效中", "已禁用", "不再进入新任务的可选问卷池。"),
        ("已禁用", "启用", "当前状态仍为已禁用", "生效中", "重新进入新任务的可选问卷池。"),
        ("任意", "删除", "未关联任何投放任务", "已删除", "从列表和可选问卷池移除。"),
        ("任意", "删除", "已关联投放任务", "不变", "提示：该问卷已关联投放任务，不允许删除。"),
    ], [1.05, 1.0, 2.15, 1.05, 2.0], 8.5)

    h(doc, 3, "5.4.7 投放任务列表")
    add_body(doc, "页面目标：查询和维护问卷投放任务。任务列表按任务 ID 倒序展示，表格内容超过页面宽度时在底部提供横向滚动条。")
    h(doc, 4, "搜索条件")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("task_name", "任务名", "否", "string", "空", "按任务名称模糊匹配。"),
        ("function_module", "功能模块", "否", "enum", "所有", "单选：所有、全局、睡眠、活动、压力、生命体征、计划、OSA、生理、血压、AI Partner。"),
        ("delivery_time_range", "投放时间", "否", "date_range", "空", "同一输入框选择开始和结束日期，选中范围高亮；按任务投放时间筛选。"),
        ("task_status", "任务状态", "否", "enum", "所有", "单选：所有、待投放、投放中、投放完成、投放失败。"),
    ], [1.25, 1.05, 0.65, 0.8, 0.8, 2.65], 8)
    h(doc, 4, "列表字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("task_id", "任务 ID", "是", "int", "系统生成", "唯一标识；列表倒序。"),
        ("task_name", "任务名称", "是", "string", "无", "展示任务名称。"),
        ("task_status", "任务状态", "是", "enum", "待投放", "待投放、投放中、投放完成、投放失败。"),
        ("function_module", "功能模块", "是", "enum", "无", "展示任务所属功能模块。"),
        ("audience", "投放人群", "是", "link", "查看", "点击进入投放人群页面。"),
        ("delivery_time", "投放时间", "是", "datetime_hour", "无", "格式 YYYY-MM-DD hh。"),
        ("completed_at", "完成时间", "否", "datetime_hour", "-", "由系统在任务结束时生成；未完成展示 -。"),
        ("questionnaire_name", "问卷名称", "是", "string", "无", "展示关联问卷名称。"),
        ("total_users", "总用户数", "是", "int", "0", "目标用户总数，千分位展示。"),
        ("actual_delivery_users", "实际投放用户数", "是", "int", "0", "实际进入投放队列的用户数。"),
        ("submitted_users", "提交用户数", "是", "int", "0", "表头含问号图标；hover：提交用户包含提交部分问卷和提交完整问卷用户。"),
        ("exposed_users", "曝光用户数", "是", "int", "0", "已曝光用户数。"),
        ("operation", "操作列", "是", "button[]", "无", "编辑、导出、删除；按钮按状态禁用。"),
    ], [1.25, 1.05, 0.65, 0.8, 0.8, 2.65], 7.8)
    h(doc, 4, "任务状态说明")
    add_table(doc, ["状态", "进入条件", "允许操作", "说明"], [
        ("待投放", "新增任务成功", "编辑、删除", "任务等待到达投放时间。"),
        ("投放中", "到达投放时间并开始下发", "查看人群", "问卷正在投放。"),
        ("投放完成", "任务下发结束，存在至少一名成功用户", "导出、查看人群", "同时存在成功和失败用户时仍为投放完成。"),
        ("投放失败", "所有用户均投放失败", "查看人群", "仅全部用户投放失败时进入。"),
    ], [1.15, 2.2, 1.6, 2.25], 8.5)
    h(doc, 4, "任务操作交互")
    add_table(doc, ["触发位置", "触发动作", "前置条件", "系统处理", "结果 / 提示"], [
        ("+新增任务", "点击", "有新增权限", "打开新增任务页面", "新增成功状态=待投放。"),
        ("编辑", "点击", "状态=待投放", "进入编辑页并锁定功能模块、问卷类型、关联问卷", "其他状态按钮禁用。"),
        ("删除", "点击", "状态=待投放", "弹出“确认删除任务吗？”", "确认删除并提示任务已删除；取消不变。"),
        ("导出", "点击", "状态=投放完成", "显示 loading 弹窗，数据准备好后触发浏览器下载", "文件名：{任务名称}_用户回复情况.xlsx；提示已加入下载队列。"),
        ("投放人群", "点击", "任务存在", "进入该任务的人群明细页", "保留投放任务导航选中状态。"),
    ], [1.15, 0.75, 1.45, 2.3, 1.55], 8)

    h(doc, 3, "5.4.8 新增 / 编辑投放任务")
    h(doc, 4, "表单字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("task_name", "任务名", "是", "string", "空", "输入框；不能为空。"),
        ("function_module", "功能模块", "是", "enum", "请选择", "全局、睡眠、活动、压力、生命体征、计划、OSA、生理、血压、AI Partner；编辑时禁用。"),
        ("questionnaire_type", "问卷类型", "条件必填", "enum", "请选择", "仅功能模块=全局时展示；选项：弹窗问卷(全局)、弹窗问卷(用研/设计)；编辑时禁用。"),
        ("questionnaire_name", "关联问卷", "是", "enum", "请选择", "只展示生效中问卷；编辑时禁用。"),
        ("delivery_time", "投放时间", "是", "datetime_hour", "空", "日期+小时；格式 YYYY-MM-DD hh；必须大于等于当前小时。"),
        ("android_versions", "Android 投放版本", "是", "array<enum>", "按规则", "多选：V3.13、V3.16；至少选择一项。"),
        ("ios_versions", "iOS 投放版本", "是", "array<enum>", "按规则", "多选：V3.13、V3.16；至少选择一项。"),
        ("audience_files", "投放人群", "否", "file[]", "未上传", "支持多个 .xlsx/.csv；单文件不超过 5M；逐行回显文件名，支持删除和下载模板。"),
    ], [1.25, 1.05, 0.65, 0.85, 0.85, 2.55], 7.9)
    h(doc, 4, "字段联动")
    add_table(doc, ["条件", "问卷类型", "关联问卷", "投放版本默认值"], [
        ("功能模块=全局", "显示且必填", "选择与问卷类型一致的生效中问卷", "用研/设计默认 V3.16；全局默认 V3.13。"),
        ("功能模块≠全局", "隐藏且清空", "可选择全部生效中的功能问卷", "睡眠/活动/压力/生命体征默认 V3.16；其他默认 V3.13。"),
        ("编辑任务", "禁用", "禁用", "回显原任务投放版本，允许修改版本字段。"),
    ], [1.55, 1.35, 2.45, 1.85], 8.5)
    add_body(doc, "当功能模块或问卷类型发生变化时，Android 与 iOS 投放版本按上述规则重新赋默认值；用户可在默认值基础上多选调整。")
    h(doc, 4, "提交与取消")
    add_table(doc, ["操作", "前置校验", "数据影响", "结果 / 提示"], [
        ("新增提交", "必填字段通过；时间不早于当前小时；Android/iOS 均至少选 1 个版本", "生成新 task_id，状态=待投放，人数均初始化为 0", "提示新增成功并返回列表。"),
        ("编辑提交", "任务仍为待投放；可编辑字段通过校验", "覆盖任务名、投放时间、投放版本、投放人群", "提示编辑成功并返回列表。"),
        ("取消", "无", "不保存修改", "返回投放任务列表。"),
    ], [1.15, 2.65, 2.15, 1.25], 8.5)

    h(doc, 3, "5.4.9 投放人群")
    add_body(doc, "页面目标：查看指定投放任务下的用户明细。页面从任务列表的投放人群链接进入，列表宽度自适应页面；超宽时展示底部横向滚动条。")
    h(doc, 4, "搜索条件")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("ringconn_id", "RingConn ID", "否", "string", "空", "按 RingConn ID 模糊查询。"),
        ("delivery_status", "投放状态", "否", "array<enum>", "所有投放状态", "多选：未开始、已关闭、已提交部分问卷、已提交全部问卷；选择后回显，超过 3 项显示已选择 N 项。"),
        ("exposure_status", "曝光状态", "否", "enum", "所有曝光状态", "单选：所有曝光状态、未曝光、已曝光。"),
    ], [1.25, 1.05, 0.65, 0.85, 0.9, 2.5], 8)
    h(doc, 4, "列表字段")
    add_table(doc, ["字段", "中文", "是否必填", "类型", "默认值", "规则 / 交互说明"], [
        ("ringconn_id", "RingConn ID", "是", "string", "无", "用户唯一标识。"),
        ("country", "国家", "是", "string", "无", "展示用户所属国家。"),
        ("registered_at", "注册时间", "是", "date", "无", "格式 YYYY-MM-DD。"),
        ("user_flag", "用户标识", "是", "enum", "无", "老用户、新用户、波动用户。"),
        ("exposure_status", "曝光状态", "是", "enum", "未曝光", "未曝光、已曝光；标签展示。"),
        ("delivery_status", "投放状态", "是", "enum", "未开始", "未开始、已关闭、已提交部分问卷、已提交全部问卷；标签展示。"),
    ], [1.25, 1.05, 0.65, 0.8, 0.85, 2.6], 8)

    h(doc, 3, "5.4.10 重复拦截与幂等规则")
    h(doc, 4, "重复拦截")
    add_table(doc, ["场景", "判重字段", "拦截规则", "提示文案"], [
        ("新增 / 编辑问卷", "待补充", "当前原型未定义最终判重组合；开发前需确认是否按问卷类型+问卷名称判重。", "待补充。"),
        ("新增 / 编辑任务", "待补充", "当前原型未定义最终判重组合；开发前需确认任务名是否允许重复。", "待补充。"),
        ("问题其他选项", "page_id + question_id + is_other", "同一问题最多一个其他选项。", "已存在时隐藏添加其他按钮。"),
    ], [1.55, 1.65, 2.8, 1.2], 8.3)
    h(doc, 4, "幂等规则")
    add_table(doc, ["场景", "幂等标识", "处理规则"], [
        ("问卷新增提交", "一次表单提交", "重复点击只生成一条问卷；成功后返回列表并禁用本次提交。"),
        ("问卷编辑提交", "questionnaire_id + 本次编辑", "同一编辑结果只应用一次。"),
        ("任务新增提交", "一次表单提交", "重复点击只生成一条任务。"),
        ("状态变更", "questionnaire_id + 目标状态", "已处于目标状态时不重复变更。"),
        ("删除", "业务 ID", "已删除数据再次删除不产生新影响。"),
        ("多语言提交", "questionnaire_id + 文件内容", "同一文件重复提交仅保留最新一次有效内容。"),
    ], [1.55, 2.25, 3.4], 8.5)

    h(doc, 1, "非功能需求")
    h(doc, 2, "财务需求")
    add_body(doc, "本期不涉及财务结算。")
    h(doc, 2, "运营或产品")
    add_bullets(doc, [
        "问卷、任务及状态变更需记录操作人、操作时间和变更结果。",
        "已禁用问卷不得出现在新增投放任务的关联问卷选项中。",
        "删除问卷前必须校验是否存在关联投放任务。",
    ])
    h(doc, 2, "UI")
    add_bullets(doc, [
        "使用现有原型规范并采用 Ant Design 配色；主要操作、成功操作和危险操作使用不同语义色。",
        "列表宽度适配页面；超出页面时展示横向滚动条，分页与表格保持 20px 间距。",
        "表单必填项展示红色星号；错误提示紧邻字段，输入框 hover/focus 显示蓝色高亮边框。",
        "表格文本默认居中；搜索条件标签右对齐；页面不出现控件重叠、越界或文本截断。",
    ])
    h(doc, 2, "研发")
    add_bullets(doc, [
        "Chrome、Edge、Safari 当前主流版本可正常使用。",
        "列表查询、状态更新、文件解析过程中需提供明确 loading 或反馈状态。",
        "富文本内容仅保留允许的加粗、下划线和字体颜色格式，避免注入风险。",
        "页面刷新或异常返回时不得生成重复业务记录。",
    ])

    h(doc, 1, "多语言")
    add_bullets(doc, [
        "弹窗问卷(全局)不开放多语言配置；其他问卷支持新增和编辑多语言文件。",
        "多语言范围：中文、繁体中文、英文、日本语、德语、西班牙语、法语、阿拉伯语、捷克语、土耳其语、波兰语、芬兰语、意大利语、匈牙利语、罗马尼亚语、韩语、立陶宛语、马来语。",
        "多语言内容覆盖网页显示名、Banner、弹窗、问题标题、问题标题备注、选项、评分描述及其他选项描述。",
        "未配置目标语言时的回退规则：待补充。",
    ])

    h(doc, 1, "埋点")
    add_table(doc, ["埋点事件", "触发时机", "属性", "说明"], [
        ("nps_questionnaire_search", "点击问卷查询", "status, questionnaire_type, has_keyword", "统计查询使用情况。"),
        ("nps_questionnaire_submit", "新增/编辑问卷提交", "mode, questionnaire_type, result", "统计提交成功率和失败原因。"),
        ("nps_questionnaire_status", "启用/禁用确认", "questionnaire_id, from_status, to_status", "统计状态变更。"),
        ("nps_questionnaire_i18n", "多语言提交", "questionnaire_id, mode, result", "统计多语言配置。"),
        ("nps_task_submit", "新增/编辑任务提交", "mode, function_module, questionnaire_type, result", "统计投放任务配置。"),
        ("nps_task_export", "点击导出", "task_id, task_status, result", "统计导出行为。"),
    ], [1.65, 1.65, 2.65, 1.25], 8.2)

    h(doc, 1, "产品项目规划（从调研、计划到转入开发）")
    add_table(doc, ["阶段", "时间", "交付物", "负责人"], [
        ("需求确认", "待补充", "PRD、原型确认", "产品"),
        ("技术评审", "待补充", "数据模型、联动规则、文件方案确认", "研发"),
        ("开发联调", "待补充", "问卷、任务、多语言和人群明细完成", "研发"),
        ("测试验收", "待补充", "测试用例、缺陷修复、产品验收", "测试 / 产品"),
        ("上线发布", "待补充", "发布计划、回滚预案、操作说明", "项目负责人"),
    ], [1.25, 1.25, 3.25, 1.45], 8.5)

    h(doc, 1, "审批与签名")
    add_table(doc, ["阶段", "签名 / 审核人", "评论要点"], [
        ("PM评审（初评）", "", ""),
        ("初评PRD", "", ""),
        ("PM评审（细评）", "", ""),
        ("细评PRD", "", ""),
        ("研发评审", "", ""),
        ("测试评审", "", ""),
    ], [2.0, 2.1, 3.1], 8.8)

    core = doc.core_properties
    core.title = "NPS管理后台需求文档"
    core.subject = "最新原型对应需求说明"
    core.author = "谢敏"
    core.keywords = "NPS, 问卷, 投放任务, PRD"
    core.comments = "按当前可见 HTML 原型范围生成"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
